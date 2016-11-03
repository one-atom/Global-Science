# coding: utf-8
import urllib2,urllib
import docx2txt
import os
import pandas as pd
import re
import string
import pdb
import csv
import xlwt
import sys
import argparse
import docx

def printPath(level, path):
    global allFileNum
    '''
    打印一个目录下的所有文件夹和文件
    '''
    # 所有文件夹，第一个字段是次目录的级别
    dirList = []
    # 所有文件
    fileList = []
    # 返回一个列表，其中包含在目录条目的名称(google翻译)
    files = os.listdir(path)
    # 先添加目录级别
    dirList.append(str(level))
    for f in files:
        if(os.path.isdir(path + '/' + f)):
            # 排除隐藏文件夹。因为隐藏文件夹过多
            if(f[0] == '.'):
                continue
            else:
                # 添加非隐藏文件夹
                dirList.append(f)
        if(os.path.isfile(path + '/' + f)):
            # 添加文件
            if(f[0] == '.'):
                continue
            else:
                fileList.append(f)
    # 当一个标志使用，文件夹列表第一个级别不打印
    i_dl = 0
    for dl in dirList:
        if(i_dl == 0):
            i_dl = i_dl + 1
        else:
            # 打印至控制台，不是第一个的目录
            print '-' * (int(dirList[0])), dl
            # 打印目录下的所有文件夹和文件，目录级别+1
            printPath((int(dirList[0]) + 1), path + '/' + dl)
    return fileList

def cut2sentence(text):
    all_sentence = []
    for key in string.split(text, '\n'):
        if key != u'':
            all_sentence.append(key)
    return all_sentence

def find_audio(path,filename):
    headers = {'User-Agent':'Mozilla/5.0 (Windows; U; Windows NT 6.1; en-US; rv:1.9.1.6) Gecko/20091201 Firefox/3.5.6'}
    text = docx2txt.process(path+'/'+filename)
    all_paragraph = cut2sentence(text)

            # english subject
    english_subject = all_paragraph[0].rstrip('/n').rstrip('/r').rstrip(' ').lstrip(' ')
    print english_subject
    words = re.findall(r'[A-z]+',english_subject)[1:-1]
    search_url = 'https://www.scientificamerican.com/search/?q=' + '+'.join(words)
   	#print search_url
    req = urllib2.Request(search_url, headers=headers)  
    content = urllib2.urlopen(req).read()
    content_url = re.search(r'\"https://www\.scientificamerican\.com/podcast/episode/(.+?)\"',content).group().rstrip('\"').lstrip('\"')
    print content_url
    req = urllib2.Request(content_url, headers=headers)  
    content = urllib2.urlopen(req).read()
    audio_url = re.search(r'\"/podcast/podcast\.mp3\?fileId=(.+?)"',content).group().rstrip('\"').lstrip('\"')
    audio_url = 'http://www.scientificamerican.com' + audio_url
    #print audio_url
    return english_subject,content_url,audio_url


def download_audio(url,path,filename):
	absolute_name = filename.split('.')[0]
	urllib.urlretrieve(url, path + '/' + absolute_name + '.mp3')

def normalize(file, url):
	doc = docx.Document(file)
	no_url = True
	for para in doc.paragraphs:
		#print para.text
		if 'https://' in para.text or 'http://' in para.text:
			no_url = False
	if no_url:
		doc.add_paragraph('')
		doc.add_paragraph(u'原文链接：',style=None)
		doc.add_paragraph(url,style=None)
	doc.save(file)

def main():
    ss_path = sys.argv[1]
    file_list = printPath(1, ss_path)
    url_file = open(ss_path + '/' +'原文链接.csv','w')
    for filename in file_list:
    	if 'docx' in filename and '~$' not in filename:
    		title, content_url, audio_url= find_audio(ss_path,filename)
    		download_audio(audio_url,ss_path, filename)
    		normalize(ss_path + '/' + filename, content_url)
    		'''
    		title = re.sub(',','',title)
    		url_file.write(filename + ',' +title.encode('utf-8')+ ','+content_url+'\n')
    		'''

if __name__ == "__main__":
    main()





